# This script creates a DOCX template with all required placeholders and standard formatting.
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

doc = Document()
doc.add_heading('{NAME}', 0)
p = doc.add_paragraph()
p.add_run('Email: {EMAIL}\n').bold = True
p.add_run('Phone: {PHONE}\n').bold = True
doc.add_paragraph('Summary:')
summary_para = doc.add_paragraph('{SUMMARY}')
doc.add_paragraph('Skills:')
skills_para = doc.add_paragraph('{PRIMARY_SKILLS}')
doc.add_paragraph('Work Experience:')
work_exp_para = doc.add_paragraph('{WORK_EXPERIENCE}')
doc.add_paragraph('Education:')
edu_para = doc.add_paragraph('{EDUCATION}')
doc.add_paragraph('Projects:')
proj_para = doc.add_paragraph('{PROJECTS}')
doc.add_paragraph('Certifications:')
cert_para = doc.add_paragraph('{CERTIFICATIONS}')

# Set font styles for all text
for para in doc.paragraphs:
    for run in para.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(11)

# Set section headings bold and larger
section_headings = ['Summary:', 'Skills:', 'Work Experience:', 'Education:', 'Projects:', 'Certifications:']
for para in doc.paragraphs:
    if para.text.strip() in section_headings:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

template_path = os.path.join('templates', 'Resume_Template_With_Logo.docx')
doc.save(template_path)
print(f"Template created at {template_path}")
