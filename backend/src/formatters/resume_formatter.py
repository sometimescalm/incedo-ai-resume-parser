from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

class ResumeFormatter:
    def _flatten_skills(self, skills):
        # Accepts dict, list, string, or None
        if skills is None:
            return []
        if isinstance(skills, dict):
            flat = []
            for v in skills.values():
                if isinstance(v, list):
                    flat.extend(v)
                elif isinstance(v, str):
                    flat.extend([s.strip() for s in v.split(',') if s.strip()])
            return flat
        elif isinstance(skills, list):
            return skills
        elif isinstance(skills, str):
            return [s.strip() for s in re.split(r'[;,]', skills) if s.strip()]
        return []

    def _add_logo_to_header_footer(self, doc, logo_path):
        if not logo_path or not os.path.exists(logo_path):
            return
        section = doc.sections[0]
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.0))
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        runf = footer_para.add_run()
        runf.add_picture(logo_path, width=Inches(1.0))
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _format_work_experience(self, work_exp):
        if work_exp is None:
            return ""
        if isinstance(work_exp, list):
            work_exp_str = ""
            for job in work_exp:
                if isinstance(job, dict):
                    work_exp_str += f"{job.get('title', '')} at {job.get('company', '')} ({job.get('duration', '')})\n"
                    details = job.get('details', job.get('description', []))
                    if isinstance(details, list):
                        for detail in details:
                            work_exp_str += f"  - {detail}\n"
                    elif isinstance(details, str):
                        work_exp_str += f"  - {details}\n"
                elif isinstance(job, str):
                    work_exp_str += job + "\n"
            return work_exp_str.strip()
        elif isinstance(work_exp, str):
            return work_exp.strip()
        return ""

    def _format_certifications(self, certs):
        if certs is None:
            return ""
        if isinstance(certs, list):
            return '\n'.join(certs)
        elif isinstance(certs, dict):
            return '\n'.join([str(v) for v in certs.values()])
        elif isinstance(certs, str):
            return certs.strip()
        return ""

    def _format_education(self, education):
        if education is None:
            return ""
        if isinstance(education, list):
            return '\n'.join(education)
        elif isinstance(education, dict):
            return '\n'.join([str(v) for v in education.values()])
        elif isinstance(education, str):
            return education.strip()
        return ""

    def format_resume_docx(self, data, template_path, logo_path="templates/Incedo_logo.PNG"):
        doc = Document(template_path)
        skills = data.get('Skills', data.get('skills', None))
        flat_skills = self._flatten_skills(skills)
        work_exp = data.get('Work Experience', data.get('work_experience', None))
        work_exp_str = self._format_work_experience(work_exp)
        certs = data.get('Certifications', data.get('certifications', None))
        certs_str = self._format_certifications(certs)
        education = data.get('Education', data.get('education', None))
        edu_str = self._format_education(education)
        replacements = {
            '{NAME}': data.get('full name', data.get('full_name', '')),
            '{EMAIL}': data.get('email id', data.get('email_id', '')),
            '{PHONE}': data.get('phone', ''),
            '{SUMMARY}': data.get('professional summary', data.get('professional_summary', '')),
            '{PRIMARY_SKILLS}': ', '.join(flat_skills),
            '{WORK_EXPERIENCE}': work_exp_str,
            '{CERTIFICATIONS}': certs_str,
            '{EDUCATION}': edu_str,
            '{LINKEDIN}': data.get('linkedin id', data.get('linkedin_id', '')),
            '{GITHUB}': data.get('github portfolio', data.get('github_portfolio', '')),
        }
        for para in doc.paragraphs:
            for key, val in replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, str(val))

        self._add_logo_to_header_footer(doc, logo_path)
        
        return doc

    def format_resume_txt(self, data, template_path):
        with open(template_path, 'r', encoding='utf-8') as f:
            template = f.read()
        skills = data.get('Skills', data.get('skills', None))
        flat_skills = self._flatten_skills(skills)
        work_exp = data.get('Work Experience', data.get('work_experience', None))
        work_exp_str = self._format_work_experience(work_exp)
        certs = data.get('Certifications', data.get('certifications', None))
        certs_str = self._format_certifications(certs)
        education = data.get('Education', data.get('education', None))
        edu_str = self._format_education(education)
        replacements = {
            '{NAME}': data.get('full name', data.get('full_name', '')),
            '{EMAIL}': data.get('email id', data.get('email_id', '')),
            '{PHONE}': data.get('phone', ''),
            '{SUMMARY}': data.get('professional summary', data.get('professional_summary', '')),
            '{PRIMARY_SKILLS}': ', '.join(flat_skills),
            '{WORK_EXPERIENCE}': work_exp_str,
            '{CERTIFICATIONS}': certs_str,
            '{EDUCATION}': edu_str,
            '{LINKEDIN}': data.get('linkedin id', data.get('linkedin_id', '')),
            '{GITHUB}': data.get('github portfolio', data.get('github_portfolio', '')),
        }
        for key, val in replacements.items():
            template = template.replace(key, str(val))
        return template